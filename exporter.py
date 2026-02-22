import io
import os
from datetime import datetime

from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT_PATH = os.path.join(os.path.dirname(__file__), 'fonts', 'ArialUnicode.ttf')


def generate_txt(results: list) -> bytes:
    lines = []
    for r in results:
        lines.append(f"--- Page {r['page']} ---")
        lines.append(r.get('translated') or '[번역 실패]')
        lines.append('')
    return '\n'.join(lines).encode('utf-8')


class _KoreanPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_font('Unicode', style='', fname=FONT_PATH)
        self.add_font('Unicode', style='B', fname=FONT_PATH)
        self.set_auto_page_break(auto=True, margin=20)

    def header(self):
        pass

    def footer(self):
        self.set_y(-15)
        self.set_font('Unicode', size=8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f'{self.page_no()}', align='C')
        self.set_text_color(0, 0, 0)


def generate_pdf(results: list, original_filename: str) -> bytes:
    pdf = _KoreanPDF()
    pdf.add_page()
    pdf.set_font('Unicode', style='B', size=20)
    pdf.ln(40)
    pdf.cell(0, 12, '번역 문서', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(6)
    pdf.set_font('Unicode', size=12)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 8, f'원본: {original_filename}', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.cell(0, 8, f'번역 일시: {datetime.now().strftime("%Y-%m-%d %H:%M")}', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.cell(0, 8, f'총 {len(results)}페이지', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.set_text_color(0, 0, 0)

    for r in results:
        pdf.add_page()
        score = r.get('similarity')
        translated = r.get('translated') or '[번역 실패]'
        low_quality = score is not None and score < 0.6

        pdf.set_font('Unicode', style='B', size=13)
        pdf.cell(0, 10, f'Page {r["page"]}', new_x='LMARGIN', new_y='NEXT')
        pdf.ln(2)

        if low_quality:
            pdf.set_fill_color(255, 243, 205)
            pdf.set_draw_color(255, 200, 0)
            pdf.set_font('Unicode', size=9)
            pdf.set_text_color(120, 80, 0)
            pdf.rect(pdf.get_x(), pdf.get_y(), 170, 8, style='FD')
            pdf.cell(170, 8, '  [!] 검토 필요 — 번역 유사도가 낮습니다', new_x='LMARGIN', new_y='NEXT')
            pdf.set_text_color(0, 0, 0)
            pdf.set_draw_color(0, 0, 0)
            pdf.ln(3)

        pdf.set_font('Unicode', size=11)
        pdf.multi_cell(0, 7, translated)
        pdf.ln(6)

        pdf.set_font('Unicode', size=8)
        pdf.set_text_color(150, 150, 150)
        pdf.cell(0, 6, f'원문 Page {r["page"]}  |  유사도: {score:.4f}' if score is not None else f'원문 Page {r["page"]}',
                 new_x='LMARGIN', new_y='NEXT')
        pdf.set_text_color(0, 0, 0)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def generate_docx(results: list, original_filename: str) -> bytes:
    doc = Document()

    title = doc.add_heading('번역 문서', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'원본: {original_filename}\n').font.size = Pt(11)
    meta.add_run(f'번역 일시: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n').font.size = Pt(11)
    meta.add_run(f'총 {len(results)}페이지').font.size = Pt(11)

    doc.add_page_break()

    for r in results:
        score = r.get('similarity')
        translated = r.get('translated') or '[번역 실패]'
        score_label = f'  (유사도: {score:.4f})' if score is not None else ''

        doc.add_heading(f'Page {r["page"]}{score_label}', level=2)

        if score is not None and score < 0.6:
            warn = doc.add_paragraph('⚠ 검토 필요 — 번역 유사도가 낮습니다')
            warn.runs[0].font.color.rgb = RGBColor(0xC0, 0x80, 0x00)
            warn.runs[0].font.size = Pt(9)

        doc.add_paragraph(translated)

        hr = doc.add_paragraph('─' * 40)
        hr.runs[0].font.size = Pt(8)
        hr.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
